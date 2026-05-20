import pandas as pd
from openai import OpenAI
from config import Config
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Groq Client
client = OpenAI(
    api_key=Config.GROQ_API_KEY,
    base_url=Config.BASE_URL
)

# The prompt enforces identifying the specific type of join (e.g., INNER JOIN)
SYSTEM_PROMPT = """
You are an expert SQL query decomposition assistant.

Your task is to analyze natural language questions and break them into:
- Intent
- Tables
- Columns
- Filters
- Joins (Always specify the type of join, e.g., INNER JOIN, LEFT OUTER JOIN, etc., followed by the join condition)

Return ONLY in this format:
Intent: ...
Tables: ...
Columns: ...
Filters: ...
Joins: ...

Do not generate SQL. Do not include any conversational filler, headers, or explanations. If no joins or filters apply, write 'None'.

=== EXAMPLE ===
User Question: "Show me all customers and their order amounts, including customers who haven't placed any orders."

Intent: Retrieve all customers alongside their order amounts, keeping unmatched customer records.
Tables: customers, orders
Columns: customers.customerNumber, customers.customerName, orders.amount
Filters: None
Joins: LEFT OUTER JOIN customers.customerNumber = orders.customerNumber
=== END OF EXAMPLE ===
"""


def decompose_query(question: str) -> str:
    """
    Convert natural language question into structured decomposition.
    """
    try:
        response = client.chat.completions.create(
            model=Config.MODEL_NAME,
            messages=[
                {
                    "role": "system",
                    "content": SYSTEM_PROMPT
                },
                {
                    "role": "user",
                    "content": f"User Question: \"{question}\""
                }
            ],
            temperature=0,
            max_tokens=250
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        return f"Error processing query: {e}"


if __name__ == "__main__":

    df = pd.read_csv("sql_questions_only.csv")
    questions = df["question"].tolist()

    doc = docx.Document()

    doc.add_heading("SQL Query Decomposition Results", level=1)

    print("\n===== QUERY DECOMPOSITION RESULTS =====\n")

    for idx, question in enumerate(questions, start=1):

        print(f"Question {idx}: {question}\n")

        result = decompose_query(question)
        print(result)
        print("\n" + "=" * 60 + "\n")

        # Write to Word
        doc.add_paragraph(f"Question {idx}: {question}", style="Heading 2")
        doc.add_paragraph(result)

    doc.save("query_decomposition_results.docx")
    print("Results saved successfully to 'query_decomposition_results.docx'")
